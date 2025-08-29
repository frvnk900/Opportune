from langchain_core.tools import tool 
from docx import Document 
import matplotlib.pyplot  as plt 
import numpy as np 
from spire.doc import  TextWatermark , WatermarkLayout , FileFormat
from spire.doc import Document as Document1
from spire.doc.common import *
from typing import Union
import os
from langchain_community.document_loaders import PyPDFLoader
from pydantic import BaseModel 
import stat


class CreateDocSchema(BaseModel):
    title:str 
    content:str 
    subtitle:str


@tool
def create_client_docx(title:str,content:str,subtitle:str) -> str:
    
    """
       Use this tool when you need to create a any document 
       args:
          title: The main title of the document based on content if not provided
          content: 
               The content to be written to the file.
               You should make the content sound more proffessional
          subtitle:
               The subtitle of the document.
               The subtitle should be based on what you are writing
          
    """
    
    # writing into document
    document = Document() 
    document.add_heading(text=f"{title.capitalize()}",level=0)
    document.add_heading(text=f"{subtitle}",level=2)
    document.add_paragraph(text=f"{content}")
    section = document.sections[0]
    footer = section.footer
    footer.add_paragraph(text='C: Property of GOBLIN.\n DO NOT EDIT')
   
     # creating file path
    BASE_DIR = os.path.abspath(os.path.dirname(__file__))
    document_name = f"myDocument{np.random.randint(100,1000)}.docx"
    UPLOAD_FOLDER = os.path.join(BASE_DIR, 'DOCX')
    save_url = os.path.join(UPLOAD_FOLDER,document_name)
    if not os.path.exists(UPLOAD_FOLDER) :
        os.makedirs(UPLOAD_FOLDER)
        
    # saving file
    document.save(save_url)
    
    
    
    def add_goblin_watermark(file:str):
    
        doc_mak = Document1()
         # creating base url
        BASE_DIR = os.path.abspath(os.path.dirname(__file__))
        UPLOAD_FOLDER = os.path.join(BASE_DIR, 'DOCX')
        save_url = os.path.join(UPLOAD_FOLDER,file)
        
        
        # appling watermark
        doc_mak.LoadFromFile(save_url)
        txtWatermark = TextWatermark()
        txtWatermark.Text = "GOBLIN DOCX."
        txtWatermark.FontSize = 80
        txtWatermark.Color = Color.get_Red()
        txtWatermark.Layout = WatermarkLayout.Diagonal
        doc_mak.Watermark = txtWatermark
        
       
 
        # saving file to base url
        doc_mak.SaveToFile(save_url)
  
    add_goblin_watermark(file=document_name)
 
    
    def remove_red_text(file,text): 
        # creating base url
            BASE_DIR = os.path.abspath(os.path.dirname(__file__))
            UPLOAD_FOLDER = os.path.join(BASE_DIR, 'DOCX')
            save_url = os.path.join(UPLOAD_FOLDER,file)
        # removing sprie.Doc licence mark
            doc_remove = Document(save_url)
            for paragraph in doc_remove.paragraphs:
                for run in paragraph.runs:
                   if text in run.text:
                      run.text = run.text.replace(text, "")
                      
        # saving changes  
            doc_remove.save(save_url)
        #  encripting file (read-only)
            os.chmod(save_url, stat.S_IRUSR | stat.S_IRGRP | stat.S_IROTH)


    text_remove = "Evaluation Warning: The document was created with Spire.Doc for Python."
    remove_red_text(file=document_name,text=text_remove)
    
#  end of tool 
 
# @tool
# def tomato_disease_management(question: str) -> str:
#     """"
#     Use the disease name from the tomato disease predictor tool to look for management ways using the the
#     disease tomato management pdf which is below based on the user's question 
#     """
#     try:
#         doc = PyPDFLoader()
#         document = doc.load()
#         embeding_vector_store=InMemoryVectorStore(Embeddings)
#         stored_docs=embeding_vector_store.add_documents(chunked_docs)
#         result=stored_docs.similarity_search(question, k=3)

#         return "\n".join([doc.page_content for doc in result])
#     except Exception as e:
#         return f"An error occurred while processing the PDF: {str(e)}"
    
 


